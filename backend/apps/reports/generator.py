from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import TYPE_CHECKING, Any

from django.conf import settings
from django.core.files.base import ContentFile
from django.template.loader import render_to_string

from core.utils.numbering import NumberGenerator

if TYPE_CHECKING:
    from .models import Report

logger = logging.getLogger(__name__)

REPORT_TEMPLATE = 'reports/report_template.html'
VERIFICATION_URL_PREFIX = getattr(
    settings, 'REPORT_VERIFICATION_URL', 'https://verify.example.com/report/',
)


def generate_report_no(commission) -> str:
    return NumberGenerator.generate(prefix='BG')


def generate_qr_verification(report_id: int) -> str:
    return f'{VERIFICATION_URL_PREFIX}{report_id}'


def _build_report_context(report: Report) -> dict:
    commission = report.commission
    items = commission.items.all() if hasattr(commission, 'items') else []

    signatures = {}
    for approval in report.approvals.select_related('user').order_by('created_at'):
        signatures[approval.role] = {
            'user': str(approval.user) if approval.user else '',
            'date': approval.created_at,
            'signature': approval.signature.url if approval.signature else None,
        }

    return {
        'report': report,
        'commission': commission,
        'items': items,
        'signatures': signatures,
        'has_cma': report.has_cma,
        'qr_code': report.qr_code,
        'conclusion': report.conclusion,
    }


def generate_report_pdf(report_id: int) -> bytes:
    from .models import Report

    report = Report.objects.select_related(
        'commission', 'compiler', 'auditor', 'approver',
    ).prefetch_related(
        'approvals__user', 'commission__items',
    ).get(pk=report_id)

    context = _build_report_context(report)

    try:
        return _generate_via_weasyprint(context)
    except ImportError:
        logger.warning('WeasyPrint not installed, falling back to placeholder')
        return _generate_placeholder(context)


def generate_report_word(report_id: int) -> bytes:
    """
    生成Word格式的检测报告
    使用python-docx库生成Word文档
    """
    from .models import Report
    
    report = Report.objects.select_related(
        'commission', 'compiler', 'auditor', 'approver',
    ).prefetch_related(
        'approvals__user', 'commission__items',
    ).get(pk=report_id)
    
    context = _build_report_context(report)
    
    try:
        return _generate_word_doc(context)
    except ImportError:
        logger.warning('python-docx not installed, falling back to placeholder')
        return _generate_placeholder(context)
    except Exception as e:
        logger.error(f'Word generation failed: {e}')
        return _generate_placeholder(context)


def _generate_via_weasyprint(context: dict) -> bytes:
    from weasyprint import HTML

    html_content = render_to_string(REPORT_TEMPLATE, context)
    pdf_buffer = io.BytesIO()
    HTML(string=html_content).write_pdf(pdf_buffer)
    return pdf_buffer.getvalue()


def _generate_word_doc(context: dict) -> bytes:
    """
    使用python-docx生成Word文档
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    
    report = context['report']
    commission = context['commission']
    items = context['items']
    signatures = context['signatures']
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('检测报告')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 报告编号
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run(f'报告编号：{report.report_no}').font.size = Pt(12)
    if report.has_cma:
        info_para.add_run('    (盖CMA章)').font.size = Pt(12)
    
    # 分隔线
    doc.add_paragraph('─' * 50)
    
    # 委托信息表格
    doc.add_paragraph('一、委托信息', style='Heading 2')
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    # 填充委托信息
    table.cell(0, 0).text = '委托单位'
    table.cell(0, 1).text = commission.client_unit
    table.cell(0, 2).text = '委托日期'
    table.cell(0, 3).text = str(commission.commission_date)
    
    table.cell(1, 0).text = '工程名称'
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(1, 1).text = str(commission.project)
    
    table.cell(2, 0).text = '施工部位'
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(2, 1).text = commission.construction_part
    
    table.cell(3, 0).text = '联系人'
    table.cell(3, 1).text = commission.client_contact or ''
    table.cell(3, 2).text = '联系电话'
    table.cell(3, 3).text = commission.client_phone or ''
    
    table.cell(4, 0).text = '见证人'
    table.cell(4, 1).text = str(commission.witness) if commission.witness else ''
    table.cell(4, 2).text = '见证取样'
    table.cell(4, 3).text = '是' if commission.is_witnessed else '否'
    
    table.cell(5, 0).text = '委托编号'
    table.cell(5, 1).text = commission.commission_no
    table.cell(5, 2).text = '委托项目数'
    table.cell(5, 3).text = str(items.count()) if hasattr(items, 'count') else str(len(items))
    
    # 检测项目表格
    doc.add_paragraph()
    doc.add_paragraph('二、检测项目', style='Heading 2')
    
    item_table = doc.add_table(rows=1, cols=6)
    item_table.style = 'Table Grid'
    
    # 表头
    headers = ['序号', '检测对象', '检测项目', '检测标准', '规格型号', '数量']
    for i, header in enumerate(headers):
        item_table.cell(0, i).text = header
        item_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
    
    # 检测项目数据
    for idx, item in enumerate(items, 1):
        row = item_table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = item.test_object
        row.cells[2].text = item.test_item
        row.cells[3].text = item.test_standard or ''
        row.cells[4].text = item.specification or ''
        row.cells[5].text = str(item.quantity)
    
    # 检测结论
    doc.add_paragraph()
    doc.add_paragraph('三、检测结论', style='Heading 2')
    doc.add_paragraph(report.conclusion or '详见检测结果')
    
    # 签章信息
    doc.add_paragraph()
    doc.add_paragraph('四、签章信息', style='Heading 2')
    
    sign_table = doc.add_table(rows=4, cols=4)
    sign_table.style = 'Table Grid'
    
    sign_table.cell(0, 0).text = '编制'
    sign_table.cell(0, 1).text = str(report.compiler) if report.compiler else ''
    sign_table.cell(0, 2).text = '编制日期'
    sign_table.cell(0, 3).text = str(report.compile_date.date()) if report.compile_date else ''
    
    sign_table.cell(1, 0).text = '审核'
    sign_table.cell(1, 1).text = str(report.auditor) if report.auditor else ''
    sign_table.cell(1, 2).text = '审核日期'
    sign_table.cell(1, 3).text = str(report.audit_date.date()) if report.audit_date else ''
    
    sign_table.cell(2, 0).text = '批准'
    sign_table.cell(2, 1).text = str(report.approver) if report.approver else ''
    sign_table.cell(2, 2).text = '批准日期'
    sign_table.cell(2, 3).text = str(report.approve_date.date()) if report.approve_date else ''
    
    sign_table.cell(3, 0).text = '发放日期'
    sign_table.cell(3, 1).text = str(report.issue_date) if report.issue_date else ''
    sign_table.cell(3, 2).text = '报告状态'
    sign_table.cell(3, 3).text = report.get_status_display()
    
    # 备注
    if report.remark:
        doc.add_paragraph()
        doc.add_paragraph('五、备注', style='Heading 2')
        doc.add_paragraph(report.remark)
    
    # 防伪二维码
    if report.qr_code:
        doc.add_paragraph()
        doc.add_paragraph(f'防伪验证：{report.qr_code}')
    
    # 页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '本报告仅对所送检样品负责，不得部分复制。'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _generate_placeholder(context: dict) -> bytes:
    report = context['report']
    text = (
        f'报告编号: {report.report_no}\n'
        f'委托编号: {context["commission"]}\n'
        f'检测结论: {report.conclusion}\n'
    )
    return text.encode('utf-8')


def generate_report_from_template(report_id: int, template_path: str = None) -> bytes:
    """
    从Word模板生成报告
    支持使用预定义的Word模板文件
    
    :param report_id: 报告ID
    :param template_path: 模板文件路径（相对于MEDIA_ROOT）
    :return: Word文档字节流
    """
    from .models import Report
    import os
    
    report = Report.objects.select_related(
        'commission', 'compiler', 'auditor', 'approver',
    ).prefetch_related(
        'approvals__user', 'commission__items',
    ).get(pk=report_id)
    
    context = _build_report_context(report)
    
    # 如果指定了模板文件，使用模板
    if template_path:
        template_full_path = os.path.join(settings.MEDIA_ROOT, template_path)
        if os.path.exists(template_full_path):
            return _generate_from_template_file(context, template_full_path)
    
    # 否则使用默认生成方式
    return _generate_word_doc(context)


def _generate_from_template_file(context: dict, template_path: str) -> bytes:
    """
    从模板文件生成Word文档
    使用Jinja2风格的占位符替换
    """
    from docx import Document
    
    doc = Document(template_path)
    
    # 替换文档中的占位符
    report = context['report']
    commission = context['commission']
    
    # 定义替换映射
    replacements = {
        '{{report_no}}': report.report_no,
        '{{commission_no}}': commission.commission_no,
        '{{client_unit}}': commission.client_unit,
        '{{project_name}}': str(commission.project),
        '{{construction_part}}': commission.construction_part,
        '{{conclusion}}': report.conclusion or '',
        '{{compile_date}}': str(report.compile_date.date()) if report.compile_date else '',
        '{{issue_date}}': str(report.issue_date) if report.issue_date else '',
    }
    
    # 替换段落中的占位符
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    
    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
